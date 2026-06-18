from pathlib import Path
import subprocess
import textwrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path("D:/Codex/home_community_battle")
DOC_DIR = ROOT / "doc"
RENDER_DIR = DOC_DIR / "rendered_teaching_docs"
DOC_DIR.mkdir(parents=True, exist_ok=True)
RENDER_DIR.mkdir(parents=True, exist_ok=True)

ACCENT = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
INK = RGBColor(23, 50, 74)
MUTED = RGBColor(80, 96, 112)
FILL = "E8EEF5"
LIGHT = "F4F6F9"
CALLOUT = "FFF7D6"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_paragraph_font(paragraph, size=11, bold=False, color=INK):
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = color


def add_para(doc, text="", style=None, size=11, bold=False, color=INK, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = color
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(7 if level == 1 else 5)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.bold = True
    run.font.size = Pt(16 if level == 1 else 13 if level == 2 else 12)
    run.font.color.rgb = ACCENT if level <= 2 else DARK
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        run.font.size = Pt(11)
        run.font.color.rgb = INK


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        run.font.size = Pt(11)
        run.font.color.rgb = INK


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = DARK
    r.font.size = Pt(11)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.line_spacing = 1.2
    set_paragraph_font(p2, size=10, color=INK)
    add_para(doc, "")


def add_table(doc, headers, rows, widths, header_fill=FILL):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header_cells = table.rows[0].cells
    set_repeat_table_header(table.rows[0])
    for idx, text in enumerate(headers):
        set_cell_shading(header_cells[idx], header_fill)
        p = header_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = DARK
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(str(text))
            r.font.size = Pt(9.5)
            r.font.color.rgb = INK
            if len(str(text)) < 10:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table


def base_doc(title, subtitle):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].font.color.rgb = INK

    footer = section.footer.paragraphs[0]
    footer.text = "愛我家鄉・社區營造大作戰"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_font(footer, size=9, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = DARK

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(14)
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(12)
    r2.font.color.rgb = MUTED
    r2.font.name = "Calibri"
    r2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    r2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    return doc


def save_doc(doc, filename):
    path = DOC_DIR / filename
    doc.save(path)
    return path


def build_lesson_plan():
    doc = base_doc("愛我家鄉：社區營造大作戰 108 課綱模式教案", "適用國小中高年級社會領域｜大螢幕四人攻防卡牌遊戲")
    add_callout(doc, "教學定位", "以社區營造放置型攻防卡牌遊戲，引導學生在收入、觀光、文化、生態、民心與效率之間做取捨，理解家鄉發展不是只有變熱鬧，而是要讓地方被看見，也能長久生活。")

    add_heading(doc, "一、課程基本資料", 1)
    add_table(doc, ["項目", "內容"], [
        ["領域/科目", "社會領域，亦可跨國語、綜合活動、資訊教育"],
        ["年級建議", "國小四年級為核心，可延伸至五下區域互動、六年級公共參與與多元文化"],
        ["節數", "2 節課，每節 40 分鐘；可擴充為 3-4 節專題課"],
        ["教學媒材", "大螢幕或互動白板、遊戲首頁、學生任務單、教師觀察紀錄表"],
        ["遊戲時間", "建議 30 分鐘一局，保留 10 分鐘整理與反思"],
    ], [2200, 7160])

    add_heading(doc, "二、108 課綱對應", 1)
    add_table(doc, ["面向", "對應精神", "本遊戲中的學習表現"], [
        ["核心素養", "自主行動、溝通互動、社會參與", "學生在有限資源下規劃社區，與同儕攻防、協商、反思公共生活。"],
        ["學習表現", "理解與思辨、態度與價值、實作與參與", "判讀地方特色、評估成本與風險、選擇防禦或陷阱、提出調整策略。"],
        ["學習內容", "家鄉特色、地方發展、公共生活、環境與文化保存", "主題卡代表地方特色；攻擊卡代表發展問題；防禦卡代表治理方法。"],
        ["議題融入", "環境教育、多元文化教育、戶外教育、資訊教育", "學生理解過度觀光、文化誤讀、生態破壞與媒體宣傳的影響。"],
    ], [1600, 2700, 5060])

    add_heading(doc, "三、學習目標", 1)
    add_bullets(doc, [
        "能說明至少一種家鄉特色如何被整理成社區主題。",
        "能判斷社區營造需要花費錢、人力、時間與技術等資源。",
        "能在觀光、收入、文化、生態、民心與效率之間做平衡決策。",
        "能解釋攻擊事件背後的社區問題，並選擇適當防禦或預防策略。",
        "能根據遊戲結局星芒圖，說出自己策略的優點與可改進方向。",
    ])

    add_heading(doc, "四、教學流程", 1)
    add_table(doc, ["階段", "時間", "教師活動", "學生活動", "評量重點"], [
        ["引起動機", "5 分", "展示首頁地圖，提問：家鄉變熱鬧一定是好事嗎？", "分享旅遊、夜市、廟會、老街或特產經驗。", "能提出地方發展的優缺點。"],
        ["規則導入", "8 分", "說明主題卡、營運卡、攻擊、防禦、陷阱、加成與商店。", "觀察卡牌圖像，猜測卡牌可能代表的社區問題。", "能分辨卡牌類型與基本用途。"],
        ["遊戲實作", "30 分", "設定玩家與 AI，提醒每回合做決策並觀察人流。", "輪流放牌、攻擊、防禦、購買卡牌與收成。", "能依狀態選擇策略，不只追求單一分數。"],
        ["結局分析", "10 分", "開啟結束分析、星芒圖與獎狀。", "閱讀自己的強項，找出一項失衡風險。", "能連結策略結果與 108 課綱能力。"],
        ["反思分享", "7 分", "引導：如果再玩一次，你會先保護哪一個指標？", "完成一句反思並分享。", "能提出具體調整策略。"],
    ], [1200, 900, 2700, 2700, 1860])

    add_heading(doc, "五、形成性與總結性評量", 1)
    add_table(doc, ["評量項目", "4 分", "3 分", "2 分", "1 分"], [
        ["認識家鄉特色", "能連結主題卡與地方故事", "能說出主題卡代表的特色", "只知道卡名，說明不足", "無法說明特色"],
        ["成本與資源判斷", "能平衡錢、人力、時間、技術", "能注意至少兩種成本", "偶爾注意成本", "只看分數或圖像"],
        ["公共生活與永續", "能兼顧民心、文化、生態", "能注意其中兩項", "只注意其中一項", "忽略公共影響"],
        ["問題解決", "能用防禦/陷阱預防危機", "遇攻擊能選擇處理", "需要提醒才處理", "無法提出處理方式"],
    ], [1800, 1890, 1890, 1890, 1890])

    add_heading(doc, "六、課後延伸", 1)
    add_bullets(doc, [
        "請學生選一張主題卡，查找真實地方故事，製作一張自己的家鄉主線卡。",
        "將班級分組，請每組提出一個家鄉營造計畫，並列出可能攻擊與防禦方法。",
        "將結局星芒圖轉成學習歷程：我做了什麼、造成什麼結果、下次怎麼調整。",
    ])
    return save_doc(doc, "愛我家鄉_108課綱模式教案.docx")


def build_player_manual():
    doc = base_doc("愛我家鄉 玩家操作手冊", "寫給小學生看的大螢幕遊戲說明")
    add_callout(doc, "一句話玩法", "你是社區營造小隊，要讓自己的家鄉更有特色、更多人來玩，也不能把居民生活、文化與環境弄壞。")

    add_heading(doc, "一、你要做什麼？", 1)
    add_bullets(doc, [
        "讓家鄉被看見：提高觀光、文化、收入。",
        "讓家鄉住得下去：照顧生態、民心、效率。",
        "保護自己的社區：遇到攻擊時用防禦或陷阱。",
        "看懂自己的策略：遊戲結束會看到星芒圖和獎狀。",
    ])

    add_heading(doc, "二、畫面上有什麼？", 1)
    add_table(doc, ["區域", "你可以做的事"], [
        ["桌牌區", "看自己和其他玩家放了哪些攻擊、防禦、加成牌；點擊可放大整理。"],
        ["手牌區", "看自己手上最多 5 張牌；點牌可放大閱讀，再決定要不要出牌。"],
        ["地方主題", "S1-S3 是重要的地方主題區，主題卡會幫你賺點數與帶來人潮。"],
        ["卡牌商店", "用點數買卡牌；主題卡很稀有，第一次購買比較便宜。"],
        ["玩家資訊", "看生命值、護甲、能量、點數、現金流、本月觀光人潮。"],
    ], [2200, 7160])

    add_heading(doc, "三、卡牌怎麼用？", 1)
    add_table(doc, ["卡牌", "放哪裡", "效果"], [
        ["主題卡", "S1-S3", "讓社區開始營運，增加點數、現金流和人潮。"],
        ["營運/加成卡", "加成區", "讓你的社區更會賺、更受歡迎或比較安全。"],
        ["攻擊卡", "指定別人", "讓對手遇到問題，例如爆量、髒亂、破財、文化誤讀。"],
        ["防禦卡", "防禦區", "先覆蓋起來，被攻擊時可能會幫你擋住危機。"],
        ["陷阱卡", "加成/陷阱區", "事前準備，等事件發生時反制。"],
    ], [1500, 1800, 6060])

    add_heading(doc, "四、輪到你時怎麼玩？", 1)
    add_numbered(doc, [
        "先看右邊玩家資訊：生命值、點數、現金流、本月人潮。",
        "點手牌，看清楚每張牌的功能。",
        "如果有主題卡，可以放到 S1-S3，讓家鄉開始營運。",
        "如果想保護自己，可以放防禦卡或陷阱卡。",
        "如果想影響對手，可以選攻擊卡，再指定要攻擊哪一位玩家。",
        "如果手牌卡住，可以在自己的回合棄牌換牌。",
        "完成後按結束回合，換下一位玩家或 AI 行動。",
    ])

    add_heading(doc, "五、小提醒", 1)
    add_bullets(doc, [
        "觀光太高可能會塞爆、髒亂或被居民抱怨。",
        "只賺錢但不顧文化，家鄉會變得沒有故事。",
        "民心太低，很多事情都推不動。",
        "防禦卡雖然是覆蓋牌，但很重要，等被攻擊時可能救你一命。",
        "不要只看誰分數最高，要看誰能平衡發展。",
    ])

    add_heading(doc, "六、遊戲結束看什麼？", 1)
    add_bullets(doc, [
        "看排名：不是只比熱鬧，也比平衡。",
        "看星芒圖：它會告訴你比較擅長哪一種能力。",
        "看獎狀：老師可以列印，記錄你這次經營家鄉的成果。",
        "想再挑戰，可以看自己哪一項比較弱，下次換策略。",
    ])
    return save_doc(doc, "愛我家鄉_玩家操作手冊_小學生版.docx")


def build_teacher_guide():
    doc = base_doc("愛我家鄉 教師教學指引", "課前準備、課中引導、課後評量與常見狀況處理")
    add_callout(doc, "教師角色", "教師不是只講規則，而是把學生的出牌選擇轉成社會領域討論：為什麼要這樣發展？誰受益？誰承擔成本？如何避免失衡？")

    add_heading(doc, "一、課前準備", 1)
    add_table(doc, ["準備項目", "建議做法"], [
        ["設備", "使用大螢幕、互動白板或投影；建議全螢幕開啟 GitHub Pages 首頁。"],
        ["座位", "四人一局最佳；若人數不足可用 AI 代玩。"],
        ["時間", "一局 30 分鐘，另留 10 分鐘結局分析與討論。"],
        ["角色分工", "可指定一位記錄員、一位規則提醒員、一位策略發言人。"],
        ["教學焦點", "本課重點不是贏牌，而是看見社區發展取捨。"],
    ], [2200, 7160])

    add_heading(doc, "二、課中提問句", 1)
    add_table(doc, ["時機", "教師可問"], [
        ["學生放主題卡", "這個地方特色是文化、產業、生態、信仰，還是觀光打卡？"],
        ["學生只衝觀光", "人變多以後，垃圾、交通、居民生活會怎麼樣？"],
        ["學生攻擊他人", "這張攻擊卡在真實社區中像什麼問題？誰需要處理？"],
        ["學生防禦成功", "這張防禦卡代表哪一種公共治理方法？"],
        ["學生陷入危機", "現在最該先救的是錢、人流、文化、生態，還是民心？為什麼？"],
        ["結局分析", "你的星芒圖哪一項最強？哪一項下次要補？"],
    ], [1800, 7560])

    add_heading(doc, "三、教師觀察紀錄", 1)
    add_table(doc, ["觀察面向", "可記錄的學生行為", "對應能力"], [
        ["認識家鄉", "能說出主題卡代表的地方特色與故事。", "家鄉特色理解"],
        ["成本判斷", "放牌前會看點數、現金流、人力或風險。", "資源與效率判斷"],
        ["公共生活", "會在攻擊或人潮過高時想到居民感受。", "公共參與與民心"],
        ["永續思考", "會避免文化、生態或民心被犧牲。", "環境與文化保存"],
        ["問題解決", "能用防禦、陷阱或換牌處理危機。", "策略調整"],
    ], [1800, 4200, 3360])

    add_heading(doc, "四、常見狀況處理", 1)
    add_table(doc, ["狀況", "教師處理方式"], [
        ["學生一直攻擊別人", "提醒攻擊也會造成自己烏雲與人流流失，引導討論衝突的成本。"],
        ["學生只買強牌", "請學生說明這張牌如何幫助社區，而不只是看分數。"],
        ["學生看不懂卡牌", "請他先點卡牌放大，再用一句話說出卡牌代表的社區問題。"],
        ["時間不夠", "可直接按結束遊戲進入分析與獎狀，保留反思時間。"],
        ["AI 太快或太亂", "把 AI 當成不同個性的同學，讓學生觀察其策略差異。"],
    ], [2200, 7160])

    add_heading(doc, "五、評量與回饋建議", 1)
    add_bullets(doc, [
        "不要只用輸贏評分，應看學生是否能說出策略理由。",
        "可用結局星芒圖當作學生自評起點，而非唯一成績。",
        "課後可請學生寫下：我讓家鄉變好的方法、造成的問題、下次的調整。",
        "獎狀可作為正向回饋，重點放在努力與理解，而不是只獎第一名。",
    ])

    add_heading(doc, "六、可延伸的跨域任務", 1)
    add_table(doc, ["任務", "做法"], [
        ["國語文", "寫一篇地方導覽詞或社區故事短文。"],
        ["綜合活動", "設計一場班級社區會議，討論觀光與居民生活的衝突。"],
        ["資訊教育", "製作一張新卡牌，包含圖片、效果、風險與防禦關係。"],
        ["戶外教育", "走訪校園附近的地方特色，拍照記錄可以營造的地方。"],
    ], [1800, 7560])
    return save_doc(doc, "愛我家鄉_教師教學指引.docx")


def main():
    outputs = [build_lesson_plan(), build_player_manual(), build_teacher_guide()]
    print("\n".join(str(path) for path in outputs))


if __name__ == "__main__":
    main()
